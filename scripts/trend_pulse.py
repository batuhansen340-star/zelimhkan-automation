# scripts/trend_pulse.py
# TrendPulse v4.0 — Kisisel Girisim Danismani + Trend Hafizasi + Bana Gore Filtresi
# 9 kaynaktan veri ceker, capraz skorlar, yasam dongusu takip eder,
# Claude ile analiz eder, DOCX rapor olusturur, Telegram'a gonderir.

import os
import sys
import json
import hashlib
import requests
import feedparser
from collections import Counter
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))
from ai_engine import ask_claude
from telegram_bot import send_telegram, send_document
from trend_memory import save_today, load_memory, get_lifecycle_summary

HEADERS = {'User-Agent': 'TrendPulse/4.0 (by /u/trendpulse_bot)'}
TODAY = datetime.now().strftime('%Y-%m-%d')
YESTERDAY = (datetime.now() - timedelta(days=1)).strftime('%Y-%m-%d')
HISTORY_FILE = os.path.join(os.path.dirname(__file__), '..', 'data', 'trend_history.json')

# Kaynak agirlik puanlari (yuksek = daha guvenilir sinyal)
SOURCE_WEIGHTS = {
    'hacker_news': 5,    # En yuksek sinyal
    'techcrunch': 4,     # Fonlama + startup haberleri
    'github': 4,         # Gelistirici trendleri
    'lobsters': 4,       # Curated tech
    'reddit': 3,         # Topluluk nabzi
    'product_hunt': 3,   # Yeni urunler
    'devto': 2,          # Gelistirici icerikleri
    'webrazzi': 3,       # Turkiye ozel
    'arxiv': 2,          # Akademik
}


# ============================================================
# VERI KAYNAKLARI (9 kaynak)
# ============================================================

def fetch_hacker_news():
    """Hacker News top stories (score > 50) — en iyi tech sinyal kaynagi"""
    print("[1/9] Hacker News cekiliyor...")
    try:
        resp = requests.get('https://hacker-news.firebaseio.com/v0/topstories.json', timeout=15)
        ids = resp.json()[:30]
        stories = []
        for item_id in ids:
            try:
                item = requests.get(f'https://hacker-news.firebaseio.com/v0/item/{item_id}.json', timeout=10).json()
                if item and item.get('score', 0) > 50:
                    stories.append({
                        'title': item.get('title', ''),
                        'url': item.get('url', f'https://news.ycombinator.com/item?id={item_id}'),
                        'score': item.get('score', 0),
                        'comments': item.get('descendants', 0),
                        '_source': 'hacker_news'
                    })
            except Exception:
                continue
        print(f"  -> {len(stories)} haber bulundu")
        return stories
    except Exception as e:
        print(f"  -> Hacker News hatasi: {e}")
        return []


def fetch_product_hunt():
    """Product Hunt RSS — yeni urun lansmanlari"""
    print("[2/9] Product Hunt cekiliyor...")
    try:
        feed = feedparser.parse('https://www.producthunt.com/feed')
        products = []
        for entry in feed.entries[:10]:
            products.append({
                'title': entry.get('title', ''),
                'url': entry.get('link', ''),
                'summary': entry.get('summary', '')[:200],
                '_source': 'product_hunt'
            })
        print(f"  -> {len(products)} urun bulundu")
        return products
    except Exception as e:
        print(f"  -> Product Hunt hatasi: {e}")
        return []


def fetch_reddit():
    """Reddit AI + Startups + Technology + SideProject (score > 30)"""
    print("[3/9] Reddit cekiliyor...")
    try:
        resp = requests.get(
            'https://reddit.com/r/artificial+startups+technology+SideProject/hot.json',
            headers=HEADERS, timeout=15
        )
        if resp.status_code != 200:
            print(f"  -> Reddit HTTP {resp.status_code}, atlaniyor")
            return []
        try:
            data = resp.json()
        except (json.JSONDecodeError, ValueError):
            print("  -> Reddit invalid JSON, atlaniyor")
            return []
        posts = []
        for child in data.get('data', {}).get('children', [])[:25]:
            post = child.get('data', {})
            if post.get('score', 0) > 30:
                posts.append({
                    'title': post.get('title', ''),
                    'url': f"https://reddit.com{post.get('permalink', '')}",
                    'score': post.get('score', 0),
                    'subreddit': post.get('subreddit', ''),
                    'comments': post.get('num_comments', 0),
                    '_source': 'reddit'
                })
        print(f"  -> {len(posts)} post bulundu")
        return posts
    except Exception as e:
        print(f"  -> Reddit hatasi: {e}")
        return []


def fetch_github_trending():
    """GitHub — son 7 gunde olusturulan en cok yildiz alan repolar"""
    print("[4/9] GitHub Trending cekiliyor...")
    try:
        week_ago = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
        gh_headers = {
            'User-Agent': 'TrendPulse/4.0',
            'Accept': 'application/vnd.github.v3+json',
        }
        gh_token = os.environ.get('GITHUB_TOKEN', '')
        if gh_token:
            gh_headers['Authorization'] = f'token {gh_token}'
        resp = requests.get(
            f'https://api.github.com/search/repositories?q=created:>{week_ago}&sort=stars&order=desc&per_page=15',
            headers=gh_headers, timeout=15
        )
        try:
            data = resp.json()
        except (json.JSONDecodeError, ValueError):
            print("  -> GitHub invalid JSON, atlaniyor")
            return []
        repos = []
        for repo in data.get('items', [])[:15]:
            repos.append({
                'name': repo.get('full_name', ''),
                'url': repo.get('html_url', ''),
                'description': (repo.get('description') or '')[:200],
                'stars': repo.get('stargazers_count', 0),
                'language': repo.get('language', 'N/A'),
                'topics': repo.get('topics', [])[:5],
                '_source': 'github'
            })
        print(f"  -> {len(repos)} repo bulundu")
        return repos
    except Exception as e:
        print(f"  -> GitHub hatasi: {e}")
        return []


def fetch_techcrunch():
    """TechCrunch RSS — fonlama haberleri ve startup ekosistemi"""
    print("[5/9] TechCrunch cekiliyor...")
    try:
        feed = feedparser.parse('https://techcrunch.com/feed/')
        articles = []
        for entry in feed.entries[:15]:
            tags = [t.get('term', '') for t in entry.get('tags', [])]
            articles.append({
                'title': entry.get('title', ''),
                'url': entry.get('link', ''),
                'summary': entry.get('summary', '')[:300],
                'tags': tags,
                'published': entry.get('published', ''),
                '_source': 'techcrunch'
            })
        print(f"  -> {len(articles)} makale bulundu")
        return articles
    except Exception as e:
        print(f"  -> TechCrunch hatasi: {e}")
        return []


def fetch_devto():
    """Dev.to API — gelistirici community trendleri"""
    print("[6/9] Dev.to cekiliyor...")
    try:
        resp = requests.get('https://dev.to/api/articles?top=1&per_page=15',
                            headers=HEADERS, timeout=15)
        try:
            data = resp.json()
        except (json.JSONDecodeError, ValueError):
            print("  -> Dev.to invalid JSON, atlaniyor")
            return []
        articles = []
        for item in data[:15]:
            articles.append({
                'title': item.get('title', ''),
                'url': item.get('url', ''),
                'description': (item.get('description') or '')[:200],
                'reactions': item.get('public_reactions_count', 0),
                'comments': item.get('comments_count', 0),
                'tags': item.get('tag_list', []),
                '_source': 'devto'
            })
        print(f"  -> {len(articles)} makale bulundu")
        return articles
    except Exception as e:
        print(f"  -> Dev.to hatasi: {e}")
        return []


def fetch_lobsters():
    """Lobste.rs — curated tech haberler, HN'den daha az gurultu"""
    print("[7/9] Lobste.rs cekiliyor...")
    try:
        resp = requests.get('https://lobste.rs/hottest.json', headers=HEADERS, timeout=15)
        try:
            data = resp.json()
        except (json.JSONDecodeError, ValueError):
            print("  -> Lobste.rs invalid JSON, atlaniyor")
            return []
        stories = []
        for item in data[:15]:
            stories.append({
                'title': item.get('title', ''),
                'url': item.get('url', '') or item.get('short_id_url', ''),
                'score': item.get('score', 0),
                'comments': item.get('comment_count', 0),
                'tags': item.get('tags', []),
                '_source': 'lobsters'
            })
        print(f"  -> {len(stories)} haber bulundu")
        return stories
    except Exception as e:
        print(f"  -> Lobste.rs hatasi: {e}")
        return []


def fetch_webrazzi():
    """Webrazzi RSS — Turkiye startup ve teknoloji ekosistemi"""
    print("[8/9] Webrazzi cekiliyor...")
    try:
        feed = feedparser.parse('https://webrazzi.com/feed/')
        articles = []
        for entry in feed.entries[:10]:
            articles.append({
                'title': entry.get('title', ''),
                'url': entry.get('link', ''),
                'summary': entry.get('summary', '')[:200],
                'published': entry.get('published', ''),
                '_source': 'webrazzi'
            })
        print(f"  -> {len(articles)} makale bulundu")
        return articles
    except Exception as e:
        print(f"  -> Webrazzi hatasi: {e}")
        return []


def fetch_arxiv():
    """ArXiv — son AI/ML arastirma makaleleri"""
    print("[9/9] ArXiv cekiliyor...")
    try:
        resp = requests.get(
            'http://export.arxiv.org/api/query?search_query=cat:cs.AI+OR+cat:cs.LG&sortBy=submittedDate&max_results=10',
            timeout=15
        )
        feed = feedparser.parse(resp.text)
        papers = []
        for entry in feed.entries:
            papers.append({
                'title': entry.get('title', '').replace('\n', ' ').strip(),
                'url': entry.get('link', ''),
                'summary': entry.get('summary', '').replace('\n', ' ')[:200],
                'authors': ', '.join([a.get('name', '') for a in entry.get('authors', [])[:3]]),
                '_source': 'arxiv'
            })
        print(f"  -> {len(papers)} makale bulundu")
        return papers
    except Exception as e:
        print(f"  -> ArXiv hatasi: {e}")
        return []


# ============================================================
# CAPRAZ KAYNAK SKORLAMA
# ============================================================

def _normalize(text):
    """Baslik normalizasyonu — kucuk harf, fazla bosluklari temizle"""
    return ' '.join(text.lower().split())


def _title_hash(text):
    """Baslik icin kisa hash olustur — benzerlik eslestirmesi icin"""
    words = _normalize(text).split()
    # Ilk 5 anlamli kelimeyi al (2+ karakter)
    key_words = [w for w in words if len(w) > 2][:5]
    return ' '.join(key_words)


def cross_source_scoring(sources):
    """Ayni konunun birden fazla kaynakta gecip gecmedigini tespit et"""
    print("\nCapraz kaynak skorlama yapiliyor...")

    all_items = []
    for source_name, items in sources.items():
        for item in items:
            item['_source'] = source_name
            item['_title_key'] = _title_hash(item.get('title', '') or item.get('name', ''))
            all_items.append(item)

    # Anahtar kelime gruplama
    keyword_groups = {}
    for item in all_items:
        words = set(_normalize(item.get('title', '') or item.get('name', '')).split())
        words = {w for w in words if len(w) > 3}  # 3+ karakter kelimeler
        for other in all_items:
            if other is item:
                continue
            other_words = set(_normalize(other.get('title', '') or other.get('name', '')).split())
            other_words = {w for w in other_words if len(w) > 3}
            # %40+ kelime eslesmesi = ayni konu
            if words and other_words:
                overlap = len(words & other_words) / min(len(words), len(other_words))
                if overlap >= 0.4:
                    key = tuple(sorted(words & other_words))
                    if key not in keyword_groups:
                        keyword_groups[key] = set()
                    keyword_groups[key].add(item['_source'])
                    keyword_groups[key].add(other['_source'])

    # Capraz kaynak sonuclari
    cross_topics = {}
    for keywords, found_sources in keyword_groups.items():
        if len(found_sources) >= 2:
            topic = ' '.join(keywords[:4])
            weight = sum(SOURCE_WEIGHTS.get(s, 1) for s in found_sources)
            cross_topics[topic] = {
                'sources': list(found_sources),
                'count': len(found_sources),
                'weight': weight
            }

    if cross_topics:
        print(f"  -> {len(cross_topics)} capraz konu tespit edildi:")
        for topic, info in sorted(cross_topics.items(), key=lambda x: x[1]['weight'], reverse=True)[:5]:
            print(f"     [{info['count']} kaynak, agirlik:{info['weight']}] {topic} ({', '.join(info['sources'])})")
    else:
        print("  -> Capraz konu bulunamadi")

    return cross_topics


# ============================================================
# TARIHSEL KARSILASTIRMA
# ============================================================

def load_history():
    """Dunku trend verilerini yukle"""
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # Son 7 gunu tut
            cutoff = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
            return {k: v for k, v in data.items() if k >= cutoff}
    except Exception as e:
        print(f"  -> Tarihce yuklenemedi: {e}")
    return {}


def save_history(today_titles):
    """Bugunun trendlerini kaydet"""
    try:
        history = load_history()
        history[TODAY] = today_titles
        os.makedirs(os.path.dirname(HISTORY_FILE), exist_ok=True)
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)
        print(f"  -> Tarihce kaydedildi ({len(today_titles)} baslik)")
    except Exception as e:
        print(f"  -> Tarihce kaydedilemedi: {e}")


def compare_with_yesterday(sources, history):
    """Dunkuyle karsilastir — yeni olan ne, devam eden ne"""
    today_titles = set()
    for items in sources.values():
        for item in items:
            title = _normalize(item.get('title', '') or item.get('name', ''))
            if title:
                today_titles.add(title)

    yesterday_titles = set()
    yesterday_data = history.get(YESTERDAY, [])
    for t in yesterday_data:
        yesterday_titles.add(_normalize(t))

    new_today = today_titles - yesterday_titles
    continuing = today_titles & yesterday_titles

    return {
        'new_count': len(new_today),
        'continuing_count': len(continuing),
        'total_today': len(today_titles),
        'today_titles': list(today_titles)
    }


# ============================================================
# CLAUDE ANALIZ — v4.0 (Lifecycle + Bana Gore)
# ============================================================

def analyze_trends(sources, cross_scores, history_comparison, lifecycle_summary):
    """Toplanan verileri Claude ile analiz et — kisisel girisim danismani v4.0"""
    print("\nClaude ile analiz ediliyor...")

    raw_data = ""
    source_labels = {
        'hacker_news': 'Hacker News (Top Stories, score > 50, AGIRLIK: 5/5)',
        'product_hunt': 'Product Hunt (Gunun Urunleri, AGIRLIK: 3/5)',
        'reddit': 'Reddit (r/artificial+startups+technology+SideProject, AGIRLIK: 3/5)',
        'github': 'GitHub Trending (Son 7 gun, en cok yildiz, AGIRLIK: 4/5)',
        'techcrunch': 'TechCrunch (Fonlama + startup haberleri, AGIRLIK: 4/5)',
        'devto': 'Dev.to (Gelistirici community, AGIRLIK: 2/5)',
        'lobsters': 'Lobste.rs (Curated tech haberler, AGIRLIK: 4/5)',
        'webrazzi': 'Webrazzi (TURKIYE startup ekosistemi, AGIRLIK: 3/5)',
        'arxiv': 'ArXiv (AI/ML arastirma, AGIRLIK: 2/5)',
    }

    for key, items in sources.items():
        label = source_labels.get(key, key)
        raw_data += f"\n## {label}:\n{json.dumps(items, ensure_ascii=False, indent=2)}\n"

    # Capraz skor bilgisi
    cross_info = ""
    if cross_scores:
        cross_info = "\n## CAPRAZ KAYNAK ANALIZI (ayni konu birden fazla kaynakta gecmis — bunlar DAHA ONEMLI):\n"
        for topic, info in sorted(cross_scores.items(), key=lambda x: x[1]['weight'], reverse=True):
            cross_info += f"- '{topic}' -> {info['count']} kaynakta ({', '.join(info['sources'])}), agirlik: {info['weight']}\n"

    # Tarihsel karsilastirma bilgisi
    history_info = ""
    if history_comparison:
        history_info = f"""
## TARIHSEL KARSILASTIRMA (dunkuyle):
- Bugun toplam {history_comparison['total_today']} benzersiz baslik
- {history_comparison['new_count']} tanesi BUGUN YENI (dun yoktu)
- {history_comparison['continuing_count']} tanesi DUN DE VARDI (devam eden trend)
YENI olanlara ONCELIK ver — devam edenleri "hala gundemde" olarak belirt.
"""

    # Yasam dongusu bilgisi (v4.0 yeni)
    lifecycle_info = ""
    if lifecycle_summary:
        lifecycle_info = f"""
## TREND YASAM DONGUSU (14 gunluk hafiza):
- Toplam takip edilen: {lifecycle_summary.get('total_tracked', 0)} trend
- Yeni (bugun ilk kez): {lifecycle_summary.get('new', 0)}
- Yukselen (2-3 gun): {lifecycle_summary.get('rising', 0)}
- Zirve (4+ gun veya 3+ kaynak): {lifecycle_summary.get('peak', 0)}
- Dusen (son 2 gun gorulmuyor): {lifecycle_summary.get('declining', 0)}
"""
        if lifecycle_summary.get('top_rising'):
            lifecycle_info += f"- EN ONEMLI yukselen trendler: {', '.join(lifecycle_summary['top_rising'][:3])}\n"
        if lifecycle_summary.get('top_peak'):
            lifecycle_info += f"- Zirvedeki trendler: {', '.join(lifecycle_summary['top_peak'][:3])}\n"
        lifecycle_info += """
Her trend icin lifecycle durumunu (new/rising/peak/declining) belirt.
RISING trendler EN DEGERLI — bunlar YARIN herkesin konusacagi sey. SIMDI HAREKET ET.
PEAK trendler artik "herkes biliyor" — farklilas veya GEC KALDIN.
DECLINING trendler "firsati kacirdin" — bunlara vakit harcama.
"""

    prompt = f"""Sen benim kisisel girisim danismanim. Adim Batuhan, Turkiye'de bankacilik BI analisti + solo gelistirici + girisimciyim. AI araclarini aktif kullaniyorum. Su an StoryPal (AI cocuk hikaye uygulamasi) ve TrendPulse (bu rapor) uzerinde calisiyorum.

Asagidaki veriler GERCEK API'lerden CEKILMISTIR. SADECE bu verilerden analiz yap. Veri UYDURMA. Emin olmadigin bir sey varsa "Yeterli veri yok" de.

SENIN GOREVIN: Her sabah bana "bugun ne yapmaliyim, nereye bakmaliyim, hangi firsat var?" soylemek. Haber bulteni YAZMA. Beni harekete gecir.

Tarih: {TODAY}

ONEMLI KURALLAR:
1. KAYNAK AGIRLIKLARI: Hacker News ve Lobste.rs'tan gelen konular DAHA ONEMLI. Webrazzi verisi Turkiye ozelinde KRITIK.
2. CAPRAZ KAYNAK: Birden fazla kaynakta gecen konular GERCEK trend — bunlara oncelik ver.
3. YENI vs DEVAM EDEN: Tarihsel karsilastirmada YENI olanlar one cikmali.
4. YASAM DONGUSU: Her trend icin lifecycle durumu (new/rising/peak/declining) belirt. RISING olanlara ONCELIK VER.
5. Her trend icin "BU SANA NE IFADE EDIYOR?" — solo gelistirici olarak bugun ne yapmaliyim?
6. Firsat varsa NET soyle: "Bu alanda Turkiye'de bosluk var, sunu yap"
7. Tool/urun cikmissa: "Bunu bugun dene, link bu" de
8. StoryPal'a uygulanabilecek bir sey varsa direkt soyle
9. Jargon YASAK — herkesin anlayacagi dilde yaz
10. Pasif ifadeler YASAK — "Sunu yap", "Bunu dene", "Bu firsati kacirma" gibi aktif ifadeler kullan
11. Turkiye pazari acisi HER trendte olsun
12. Webrazzi verisini Turkiye Kosesi bolumu icin ozellikle kullan
13. BANA GORE FILTRESI: Her trend icin "for_me" skoru ver (0-3):
    - solo_dev: Solo gelistirici olarak yapabilir miyim? (0=hayir, 1=zor, 2=uygun, 3=mukemmel)
    - low_budget: Dusuk butceyle ($0-100/ay) mumkun mu? (0=imkansiz, 1=pahali, 2=uygun, 3=ucretsiz)
    - turkey_market: Turkiye pazarinda karsiligi var mi? (0=yok, 1=az, 2=var, 3=buyuk firsat)
    - total: Toplam skor (0-9) — 7+ = KESINLIKLE BAK, 4-6 = deger, 0-3 = atla

Gorevlerin:
1. MANSET: Bugunun en onemli 1 cumlelik ozeti — gazete manseti gibi, vurucu
2. BUGUN NE YAP: BUGUN uygulanabilecek 3 somut aksiyon (link dahil). "SIMDI YAP" formatinda.
3. TOP 5 TREND: Her trend icin:
   - Kisa baslik (max 8 kelime) + tek emoji
   - 2 cumle neden onemli (jargonsuz)
   - "Sana ne:" 1 cumle aksiyon
   - "Turkiye:" 1 cumle
   - Etki puani (1-10)
   - Yeni mi yoksa devam eden trend mi? (new/continuing)
   - Kac kaynakta gecti?
   - Yasam dongusu: new/rising/peak/declining
   - for_me skoru
4. FIRSAT RADAR: En guclu 1 uygulama firsati — detayli MVP plani
5. AI SPOTLIGHT: 1 AI tool/model — "bunu su isine kullanabilirsin" formatinda
6. PARA NEREYE AKIYOR: TechCrunch verisine dayanarak. Yoksa "veri yok" de.
7. TURKIYE KOSESI: Webrazzi verisine dayanarak Turkiye ekosistemi yorumu. Webrazzi verisi bossa genel Turkiye yorumu yap.
8. STORYPAL ICIN: Uygulanabilecek 1 sey varsa yaz. Yoksa null.

Format: Sadece JSON dondur:
{{
  "date": "{TODAY}",
  "headline": "Vurucu manset, max 15 kelime",
  "today_actions": [
    {{"action": "Sunu yap", "link": "url", "why": "Cunku..."}},
    {{"action": "Bunu dene", "link": "url", "why": "Cunku..."}},
    {{"action": "Suna bak", "link": "url", "why": "Cunku..."}}
  ],
  "executive_summary": "3 cumle, aktif dil",
  "top_trends": [
    {{
      "title": "Max 8 kelime",
      "emoji": "tek emoji",
      "impact_score": 8,
      "why": "2 cumle, jargonsuz",
      "action_for_you": "1 cumle aksiyon",
      "turkey": "1 cumle Turkiye acisi",
      "sources": ["kaynak1", "kaynak2"],
      "source_count": 2,
      "is_new": true,
      "lifecycle": "new|rising|peak|declining",
      "category": "AI|Startup|Altyapi|Yaratici|Arastirma",
      "for_me": {{
        "solo_dev": 2,
        "low_budget": 3,
        "turkey_market": 1,
        "total": 6,
        "verdict": "1 cumle neden sana uygun/uygun degil"
      }}
    }}
  ],
  "opportunity": {{
    "name": "Fikir adi",
    "one_liner": "1 cumle",
    "who_buys": "Kime satilir",
    "turkey_competitor": "Rakip var mi",
    "mvp_weeks": "Kac hafta",
    "mvp_stack": "Teknolojiler",
    "mvp_cost": "Maliyet",
    "free_hook": "Ucretsiz ne verilir",
    "paid_product": "Ucretli ne satilir",
    "why_now": "Bu firsati kacirma cunku...",
    "for_me": {{
      "solo_dev": 3,
      "low_budget": 2,
      "turkey_market": 2,
      "total": 7
    }}
  }},
  "ai_tool": {{
    "name": "Adi",
    "what": "Ne yapiyor",
    "use_case": "Senin icin ne ifade ediyor",
    "link": "url"
  }},
  "money_flow": "1-2 cumle",
  "turkey_corner": "Webrazzi verisine dayali 2-3 cumle Turkiye yorumu",
  "storypal_tip": "1 cumle veya null",
  "data_quality": {{
    "total_sources": 9,
    "active_sources": 0,
    "total_items": 0,
    "cross_source_topics": 0,
    "new_today": 0
  }},
  "sources": ["url1", "url2"]
}}
Turkce yaz. Kisa, net, aksiyon odakli. Beni harekete gecir.
{cross_info}
{history_info}
{lifecycle_info}
VERILER:
{raw_data}
"""

    result = ask_claude(prompt, json_mode=True)
    print("  -> Analiz tamamlandi")
    return result


# ============================================================
# DOCX RAPOR OLUSTURMA — v4.0
# ============================================================

def _set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex
    })
    shading.append(shading_elem)


def _for_me_badge(total_score):
    """Bana Gore skoru icin renk ve emoji"""
    try:
        s = int(total_score)
    except (ValueError, TypeError):
        s = 0
    if s >= 7:
        return '\U0001F7E2', 'SANA GORE', '27AE60'  # yesil
    elif s >= 4:
        return '\U0001F7E1', 'DEGER', 'F39C12'  # sari
    elif s >= 2:
        return '\U0001F7E0', 'BAKILABILIR', 'E67E22'  # turuncu
    else:
        return '\U0001F534', 'ATLA', 'E74C3C'  # kirmizi


def _lifecycle_badge(lifecycle):
    """Yasam dongusu icin emoji ve renk"""
    badges = {
        'new': ('\U0001F195', 'YENI', 'E74C3C'),
        'rising': ('\U0001F4C8', 'YUKSELIYOR', '27AE60'),
        'peak': ('\U0001F525', 'ZIRVEDE', 'F39C12'),
        'declining': ('\U0001F4C9', 'DUSUYOR', '95A5A6'),
    }
    return badges.get(lifecycle, ('\u2753', lifecycle or '?', '7F8C8D'))


def create_docx_report(analysis, lifecycle_summary):
    """Kisisel girisim danismani DOCX raporu — v4.0"""
    print("\nDOCX rapor olusturuluyor...")

    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)

    BLUE_DARK = RGBColor(0x1B, 0x4F, 0x72)
    BLUE_MED = RGBColor(0x2E, 0x86, 0xC1)
    ORANGE = RGBColor(0xF3, 0x9C, 0x12)
    GREEN = RGBColor(0x27, 0xAE, 0x60)
    GRAY = RGBColor(0x7F, 0x8C, 0x8D)
    GRAY_LIGHT = RGBColor(0x95, 0xA5, 0xA6)
    RED = RGBColor(0xE7, 0x4C, 0x3C)

    def run(para, text, color=None, bold=False, italic=False, size=None):
        r = para.add_run(text)
        if color: r.font.color.rgb = color
        if bold: r.bold = True
        if italic: r.italic = True
        if size: r.font.size = Pt(size)
        return r

    def h(text, level=1):
        hd = doc.add_heading(text, level=level)
        for r in hd.runs: r.font.color.rgb = BLUE_DARK
        return hd

    def sep():
        p = doc.add_paragraph()
        r = p.add_run('_' * 65)
        r.font.color.rgb = GRAY_LIGHT
        r.font.size = Pt(6)

    def stars(score):
        try: s = int(score)
        except (ValueError, TypeError): s = 5
        return '\u2B50' * s + '\u2606' * (10 - s)

    # === KAPAK ===
    doc.add_paragraph()
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(cover, '\U0001F4C8 TrendPulse v4.0', color=BLUE_DARK, bold=True, size=28)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(dp, analysis.get('date', TODAY), color=GRAY, size=12)

    hl = analysis.get('headline', analysis.get('daily_headline', ''))
    if hl:
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(hp, hl, italic=True, color=BLUE_MED, size=18)

    # Veri kalitesi badge
    dq = analysis.get('data_quality', {})
    if dq:
        badge_p = doc.add_paragraph()
        badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(badge_p, f"9 kaynak | {dq.get('total_items', '?')} veri noktasi | {dq.get('cross_source_topics', '?')} capraz konu | {dq.get('new_today', '?')} yeni bugun",
            color=GRAY_LIGHT, size=9)

    # Yasam dongusu ozeti
    if lifecycle_summary:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lc_text = (
            f"\U0001F195 {lifecycle_summary.get('new', 0)} yeni | "
            f"\U0001F4C8 {lifecycle_summary.get('rising', 0)} yukseliyor | "
            f"\U0001F525 {lifecycle_summary.get('peak', 0)} zirvede | "
            f"\U0001F4C9 {lifecycle_summary.get('declining', 0)} dusuyor"
        )
        run(lp, lc_text, color=GRAY, size=9)

    sep()
    doc.add_page_break()

    # === BUGUN NE YAP ===
    h('\U0001F3AF BUGUN NE YAP', level=1)
    actions = analysis.get('today_actions', [])
    if actions:
        for i, act in enumerate(actions[:3], 1):
            action_text = act.get('action', '') if isinstance(act, dict) else str(act)
            why_text = act.get('why', '') if isinstance(act, dict) else ''
            link_text = act.get('link', '') if isinstance(act, dict) else ''
            t = doc.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = t.rows[0].cells[0]
            _set_cell_shading(cell, 'FEF9E7')
            cp = cell.paragraphs[0]
            run(cp, f'  {i}. ', bold=True, color=ORANGE, size=13)
            run(cp, action_text, bold=True, color=BLUE_DARK, size=12)
            if why_text:
                wp = cell.add_paragraph()
                run(wp, f'     {why_text}', color=GRAY, italic=True, size=9)
            if link_text:
                lp2 = cell.add_paragraph()
                run(lp2, f'     {link_text}', color=BLUE_MED, size=8)
            doc.add_paragraph()
    else:
        ta = analysis.get('today_action', '')
        if ta:
            p = doc.add_paragraph()
            run(p, ta, bold=True, color=BLUE_DARK, size=13)
    sep()

    # === YONETICI OZETI ===
    h('\u2615 Yonetici Ozeti', level=1)
    st = doc.add_table(rows=1, cols=1)
    sc = st.rows[0].cells[0]
    _set_cell_shading(sc, 'EBF5FB')
    sp = sc.paragraphs[0]
    run(sp, analysis.get('executive_summary', ''), size=12)
    np = sc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(np, '\u2615 2 dk okuma | 9 kaynak | capraz skorlanmis | yasam dongusu takipli', color=GRAY_LIGHT, italic=True, size=8)
    doc.add_paragraph()

    # === TOP 5 TREND ===
    h('\U0001F525 TOP 5 TREND', level=1)
    trends = analysis.get('top_trends', [])
    for i, trend in enumerate(trends[:5], 1):
        emoji = trend.get('emoji', '\U0001F525')
        title = trend.get('title', '')
        score = trend.get('impact_score', 5)
        is_new = trend.get('is_new', True)
        src_count = trend.get('source_count', 1)
        lifecycle = trend.get('lifecycle', 'new')
        for_me = trend.get('for_me', {})
        for_me_total = for_me.get('total', 0) if isinstance(for_me, dict) else 0

        th = doc.add_paragraph()
        run(th, f'{emoji} {i}. {title}', bold=True, color=BLUE_DARK, size=13)

        # Lifecycle badge
        lc_emoji, lc_text, lc_color = _lifecycle_badge(lifecycle)
        run(th, f'  {lc_emoji} {lc_text}', bold=True, color=RGBColor(
            int(lc_color[0:2], 16), int(lc_color[2:4], 16), int(lc_color[4:6], 16)
        ), size=9)

        # Source count badge
        if src_count > 1:
            run(th, f'  [{src_count} kaynak]', color=GREEN, size=9)

        # Bana Gore badge
        fm_emoji, fm_text, fm_color = _for_me_badge(for_me_total)
        run(th, f'  {fm_emoji} {fm_text} ({for_me_total}/9)', bold=True, color=RGBColor(
            int(fm_color[0:2], 16), int(fm_color[2:4], 16), int(fm_color[4:6], 16)
        ), size=9)

        sp2 = doc.add_paragraph()
        run(sp2, f'Etki: {stars(score)}', size=10)

        why_text = trend.get('why', trend.get('why_important', ''))
        if why_text:
            wp = doc.add_paragraph()
            run(wp, why_text, size=11)

        at = trend.get('action_for_you', trend.get('so_what', ''))
        if at:
            ap = doc.add_paragraph()
            run(ap, '\U0001F3AF Sana ne: ', bold=True, color=BLUE_DARK, size=11)
            run(ap, at, color=BLUE_DARK, size=11)

        tt = trend.get('turkey', trend.get('turkey_angle', ''))
        if tt:
            tp = doc.add_paragraph()
            run(tp, '\U0001F1F9\U0001F1F7 Turkiye: ', bold=True, size=10)
            run(tp, tt, italic=True, color=GRAY, size=10)

        # Bana Gore detay
        if isinstance(for_me, dict) and for_me.get('verdict'):
            fmp = doc.add_paragraph()
            run(fmp, f'{fm_emoji} Bana Gore: ', bold=True, size=10)
            run(fmp, for_me['verdict'], italic=True, size=10)
            # Skor detaylari
            sd_text = f" [Solo:{for_me.get('solo_dev', '?')} Butce:{for_me.get('low_budget', '?')} TR:{for_me.get('turkey_market', '?')}]"
            run(fmp, sd_text, color=GRAY_LIGHT, size=8)

        ts = trend.get('sources', [])
        if ts:
            srcp = doc.add_paragraph()
            run(srcp, 'Kaynaklar: ' + ', '.join(str(s) for s in ts), color=GRAY_LIGHT, size=8)

        if i < len(trends[:5]):
            sep()
    doc.add_paragraph()

    # === FIRSAT RADAR ===
    h('\U0001F4A1 FIRSAT RADAR', level=1)
    opp = analysis.get('opportunity', analysis.get('opportunity_radar', {}))
    if opp:
        on = opp.get('name', opp.get('idea', ''))
        ol = opp.get('one_liner', '')
        opp_for_me = opp.get('for_me', {})
        opp_total = opp_for_me.get('total', 0) if isinstance(opp_for_me, dict) else 0
        if on:
            onp = doc.add_paragraph()
            run(onp, on, bold=True, color=BLUE_DARK, size=14)
            if ol:
                run(onp, f' \u2014 {ol}', italic=True, color=GRAY, size=11)
            # Bana Gore badge for opportunity
            fm_emoji, fm_text, fm_color = _for_me_badge(opp_total)
            run(onp, f'  {fm_emoji} {fm_text} ({opp_total}/9)', bold=True, color=RGBColor(
                int(fm_color[0:2], 16), int(fm_color[2:4], 16), int(fm_color[4:6], 16)
            ), size=10)

        ot = doc.add_table(rows=6, cols=2)
        ot.style = 'Light Grid Accent 1'
        ot.alignment = WD_TABLE_ALIGNMENT.CENTER
        rd = [
            ('Kime Satilir', opp.get('who_buys', opp.get('target_market', ''))),
            ('Turkiye Rakip', opp.get('turkey_competitor', opp.get('competitors', ''))),
            ('MVP Suresi', opp.get('mvp_weeks', opp.get('mvp_time', ''))),
            ('Tech Stack', opp.get('mvp_stack', '')),
            ('Maliyet', opp.get('mvp_cost', '')),
            ('Turkiye Potansiyeli', opp.get('turkey_potential', '')),
        ]
        for i, (label, value) in enumerate(rd):
            ot.rows[i].cells[0].text = label
            ot.rows[i].cells[1].text = str(value) if value else ''
        doc.add_paragraph()

        fh = opp.get('free_hook', opp.get('zelimkhan_hook', ''))
        pp = opp.get('paid_product', '')
        if fh:
            fhp = doc.add_paragraph()
            run(fhp, '\U0001F3A3 Zelimkhan Hook: ', bold=True, color=ORANGE, size=12)
            if pp:
                run(fhp, f'Ucretsiz \u2192 {fh} | Ucretli \u2192 {pp}', color=ORANGE, size=11)
            else:
                run(fhp, fh, color=ORANGE, size=11)

        wn = opp.get('why_now', '')
        if wn:
            wnp = doc.add_paragraph()
            run(wnp, f'\u26A1 {wn}', bold=True, color=BLUE_MED, size=11)
    doc.add_paragraph()

    # === AI SPOTLIGHT ===
    h('\U0001F916 AI SPOTLIGHT', level=1)
    ai = analysis.get('ai_tool', analysis.get('ai_spotlight', {}))
    if ai:
        an = ai.get('name', ai.get('title', ''))
        aw = ai.get('what', ai.get('detail', ''))
        au = ai.get('use_case', ai.get('practical_use', ''))
        al = ai.get('link', '')
        if an:
            anp = doc.add_paragraph()
            run(anp, an, bold=True, color=BLUE_DARK, size=13)
        if aw:
            awp = doc.add_paragraph()
            run(awp, aw, size=11)
        if au:
            aup = doc.add_paragraph()
            run(aup, '\U0001F4A1 Senin icin: ', bold=True, color=BLUE_DARK, size=11)
            run(aup, au, color=BLUE_DARK, size=11)
        if al:
            alp = doc.add_paragraph()
            run(alp, al, color=BLUE_MED, size=9)
    doc.add_paragraph()

    # === PARA NEREYE AKIYOR ===
    h('\U0001F4B0 PARA NEREYE AKIYOR', level=1)
    mf = analysis.get('money_flow', '')
    if isinstance(mf, dict):
        mt = f"{mf.get('title', '')} \u2014 {mf.get('detail', '')}"
    else:
        mt = str(mf) if mf else 'Bu hafta one cikan fonlama haberi yok'
    mp = doc.add_paragraph()
    run(mp, mt, size=11)
    doc.add_paragraph()

    # === TURKIYE KOSESI ===
    h('\U0001F1F9\U0001F1F7 TURKIYE KOSESI', level=1)
    tc = analysis.get('turkey_corner', '')
    if tc:
        tcp = doc.add_paragraph()
        run(tcp, str(tc), size=11)
    doc.add_paragraph()

    # === STORYPAL IPUCU ===
    spt = analysis.get('storypal_tip')
    if spt:
        h('\U0001F4F1 STORYPAL IPUCU', level=1)
        stt = doc.add_table(rows=1, cols=1)
        stc = stt.rows[0].cells[0]
        _set_cell_shading(stc, 'EAFAF1')
        stp = stc.paragraphs[0]
        run(stp, str(spt), color=GREEN, size=11)
        doc.add_paragraph()

    # === KAYNAKLAR ===
    h('\U0001F517 KAYNAKLAR', level=1)
    links = analysis.get('sources', analysis.get('source_links', []))
    for i, link in enumerate(links, 1):
        lp3 = doc.add_paragraph()
        run(lp3, f'{i}. {link}', color=BLUE_MED, size=9)

    # Footer
    doc.add_paragraph()
    sep()
    ft = doc.add_paragraph()
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(ft, f'TrendPulse v4.0 by Zelimkhan Automation | 9 kaynak | capraz skorlama | yasam dongusu | bana gore filtresi | {TODAY}', color=GRAY_LIGHT, size=8)

    filename = f'TrendPulse_{TODAY}.docx'
    doc.save(filename)
    print(f"  -> Rapor kaydedildi: {filename}")
    return filename


# ============================================================
# TELEGRAM GONDERIM — v4.0
# ============================================================

def send_report(analysis, docx_path, lifecycle_summary):
    """Raporu Telegram'a gonder — v4.0"""
    print("\nTelegram'a gonderiliyor...")

    headline = analysis.get('headline', analysis.get('daily_headline', ''))

    actions = analysis.get('today_actions', [])
    actions_text = ""
    for i, act in enumerate(actions[:3], 1):
        if isinstance(act, dict):
            actions_text += f"\n{i}. {act.get('action', '')}"
        else:
            actions_text += f"\n{i}. {act}"
    if not actions_text:
        ta = analysis.get('today_action', '')
        if ta: actions_text = f"\n1. {ta}"

    trends_text = ""
    for i, t in enumerate(analysis.get('top_trends', [])[:5], 1):
        emoji = t.get('emoji', '')
        title = t.get('title', '')
        lifecycle = t.get('lifecycle', 'new')
        for_me = t.get('for_me', {})
        for_me_total = for_me.get('total', 0) if isinstance(for_me, dict) else 0

        # Lifecycle emoji
        lc_map = {'new': '\U0001F195', 'rising': '\U0001F4C8', 'peak': '\U0001F525', 'declining': '\U0001F4C9'}
        lc_emoji = lc_map.get(lifecycle, '')

        # Bana Gore emoji
        fm_emoji, _, _ = _for_me_badge(for_me_total)

        sc = t.get('source_count', 1)
        multi = f' [{sc}x]' if sc > 1 else ''
        trends_text += f"\n{i}. {emoji} {title} {lc_emoji}{multi} {fm_emoji}{for_me_total}/9"

    opp = analysis.get('opportunity', analysis.get('opportunity_radar', {}))
    opp_name = opp.get('name', opp.get('idea', ''))
    opp_one = opp.get('one_liner', '')
    opp_mvp = opp.get('mvp_weeks', opp.get('mvp_time', ''))
    opp_cost = opp.get('mvp_cost', '')
    opp_for_me = opp.get('for_me', {})
    opp_total = opp_for_me.get('total', 0) if isinstance(opp_for_me, dict) else 0
    opp_fm_emoji, _, _ = _for_me_badge(opp_total)

    ai = analysis.get('ai_tool', analysis.get('ai_spotlight', {}))
    ai_name = ai.get('name', ai.get('title', ''))
    ai_use = ai.get('use_case', ai.get('practical_use', ''))

    dq = analysis.get('data_quality', {})

    # Lifecycle ozet
    lc_text = ""
    if lifecycle_summary:
        lc_text = (
            f"\n\U0001F195 {lifecycle_summary.get('new', 0)} yeni | "
            f"\U0001F4C8 {lifecycle_summary.get('rising', 0)} yukseliyor | "
            f"\U0001F525 {lifecycle_summary.get('peak', 0)} zirvede | "
            f"\U0001F4C9 {lifecycle_summary.get('declining', 0)} dusuyor"
        )

    message = f"""\U0001F4C8 *TrendPulse v4.0* \u2014 {TODAY}

\U0001F5DE *{headline}*

\U0001F3AF *BUGUN NE YAP:*{actions_text}

\U0001F525 *Top 5:*{trends_text}

\U0001F4A1 *Firsat:* {opp_name} \u2014 {opp_one}
\u23F1 MVP: {opp_mvp} | \U0001F4B0 {opp_cost} | {opp_fm_emoji} {opp_total}/9

\U0001F916 *Bugun dene:* {ai_name} \u2192 {ai_use}

\U0001F4CA _{dq.get('total_items', '?')} veri | 9 kaynak | {dq.get('cross_source_topics', '?')} capraz konu_{lc_text}

_Detayli rapor_ \u2B07\uFE0F"""

    send_telegram(message)
    send_document(docx_path, caption=f"TrendPulse v4.0 - {TODAY}")
    print("  -> Telegram gonderimi tamamlandi")


# ============================================================
# MAIN — v4.0
# ============================================================

def main():
    print(f"{'='*50}")
    print(f"  TrendPulse v4.0 - Kisisel Girisim Danismani")
    print(f"  {TODAY} | 9 kaynak | capraz skorlama | yasam dongusu | bana gore")
    print(f"{'='*50}\n")

    # 1. Veri topla (9 kaynak)
    sources = {
        'hacker_news': fetch_hacker_news(),
        'product_hunt': fetch_product_hunt(),
        'reddit': fetch_reddit(),
        'github': fetch_github_trending(),
        'techcrunch': fetch_techcrunch(),
        'devto': fetch_devto(),
        'lobsters': fetch_lobsters(),
        'webrazzi': fetch_webrazzi(),
        'arxiv': fetch_arxiv(),
    }

    total = sum(len(v) for v in sources.values())
    active = sum(1 for v in sources.values() if v)
    print(f"\nToplam {total} veri noktasi toplandi ({active}/9 kaynak aktif).")
    for name, data in sources.items():
        status = '\u2705' if data else '\u274C'
        weight = SOURCE_WEIGHTS.get(name, 1)
        print(f"  {status} {name}: {len(data)} veri (agirlik: {weight}/5)")

    if total == 0:
        print("HATA: Hicbir kaynaktan veri alinamadi!")
        send_telegram("*TrendPulse HATA*\nHicbir kaynaktan veri alinamadi.")
        sys.exit(1)

    # 2. Capraz kaynak skorlama
    cross_scores = cross_source_scoring(sources)

    # 3. Tarihsel karsilastirma
    history = load_history()
    history_comparison = compare_with_yesterday(sources, history)
    print(f"\nTarihsel karsilastirma:")
    print(f"  Yeni bugun: {history_comparison['new_count']}")
    print(f"  Devam eden: {history_comparison['continuing_count']}")

    # 4. Trend hafizasini guncelle ve yasam dongusu analizi (v4.0 yeni)
    print("\nTrend hafizasi guncelleniyor...")
    memory = save_today(sources, TODAY)
    lifecycle_summary = get_lifecycle_summary(memory)
    print(f"  Yasam dongusu:")
    print(f"  \U0001F195 Yeni: {lifecycle_summary.get('new', 0)}")
    print(f"  \U0001F4C8 Yukseliyor: {lifecycle_summary.get('rising', 0)}")
    print(f"  \U0001F525 Zirvede: {lifecycle_summary.get('peak', 0)}")
    print(f"  \U0001F4C9 Dusuyor: {lifecycle_summary.get('declining', 0)}")
    print(f"  Toplam takip: {lifecycle_summary.get('total_tracked', 0)}")

    # 5. Claude ile analiz (lifecycle ve bana gore bilgisi dahil)
    analysis = analyze_trends(sources, cross_scores, history_comparison, lifecycle_summary)

    # Veri kalitesi bilgisi ekle
    if 'data_quality' not in analysis or not isinstance(analysis.get('data_quality'), dict):
        analysis['data_quality'] = {}
    analysis['data_quality'].update({
        'total_sources': 9,
        'active_sources': active,
        'total_items': total,
        'cross_source_topics': len(cross_scores),
        'new_today': history_comparison['new_count']
    })

    # 6. DOCX rapor olustur (lifecycle bilgisi dahil)
    docx_path = create_docx_report(analysis, lifecycle_summary)

    # 7. Telegram'a gonder (lifecycle bilgisi dahil)
    send_report(analysis, docx_path, lifecycle_summary)

    # 8. Tarihceyi kaydet
    save_history(history_comparison.get('today_titles', []))

    # 9. Temizlik
    if os.path.exists(docx_path):
        os.remove(docx_path)

    print(f"\n{'='*50}")
    print("  TrendPulse v4.0 tamamlandi!")
    print(f"{'='*50}")


if __name__ == '__main__':
    main()
